"""
Generate the Veritas AI Construction Platform business case Word document.
Run once from project root: python scripts/generate_business_case.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = PROJECT_ROOT / "docs" / "assets" / "knightroad_veritas_logo.png"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "Veritas_AI_Construction_Business_Case.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
CYAN = RGBColor(0x00, 0x9E, 0xC4)


def set_run_font(run, size_pt=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def add_body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3A5C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- Logo & title block ---
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.8))
    else:
        add_body(doc, "[Company logo]", italic=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("KnightRoad Veritas AI Platforms")
    set_run_font(t_run, size_pt=14, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Business Case — Veritas AI Construction Platform")
    set_run_font(s_run, size_pt=12, color=CYAN)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run("Prepared for Executive Review  |  June 2026  |  Confidential")
    set_run_font(m_run, size_pt=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    add_horizontal_rule(doc)

    # --- EXECUTIVE BRIEF (one page) ---
    add_heading(doc, "Executive Brief", level=1)

    add_body(
        doc,
        "KnightRoad Veritas AI Platforms delivers purpose-built AI platforms that turn "
        "operational pain points into measurable outcomes. Veritas AI Construction Platform "
        "is our flagship construction solution — a single command center that helps "
        "organizations stand up projects faster, keep safety and schedules aligned, train "
        "workforces before they reach the site, and give leaders real-time visibility without "
        "juggling disconnected tools.",
        space_after=10,
    )

    add_heading(doc, "The Opportunity", level=2)
    add_body(
        doc,
        "Construction teams lose time and money to fragmented systems: one tool for scheduling, "
        "another for safety, spreadsheets for assets, and training tracked separately. "
        "Compliance gaps, schedule drift, and untrained workers on site drive cost overruns "
        "and liability. Vocational programs face the same problem while also teaching students "
        "how real projects run.",
        space_after=8,
    )

    add_heading(doc, "What Veritas Delivers", level=2)
    add_bullet(doc, "Guided 10-step project setup that auto-builds schedules, safety rules, site zones, and training assignments from proven building templates.")
    add_bullet(doc, "Executive dashboard with live progress, budget tracking, safety alerts, tasks, team roster, and 3D building models in one view.")
    add_bullet(doc, "Design Studio — in-browser 2D floor plan drawing and instant 3D model preview, with IFC4 export and Speckle collaboration — no CAD software required.")
    add_bullet(doc, "Integrated safety monitoring with zone-based alerts and acknowledgement workflows tied to OSHA and institutional standards.")
    add_bullet(doc, "Resource planning with Kanban boards and critical-path scheduling, plus asset and personnel tracking.")
    add_bullet(doc, "VR training hub that assigns role-based competency modules at project creation and tracks completion.")
    add_bullet(doc, "Eight pre-built building templates (residential, commercial, healthcare, vocational, civil, and more) so new projects launch in hours, not weeks.")

    add_heading(doc, "Proof Point", level=2)
    add_body(
        doc,
        "The platform is deployed as a working product (v1.0) for BTVI Vocational School's "
        "Construction Department — managing a BSD$1.5M vocational center build in Nassau. "
        "It demonstrates KnightRoad's ability to design, build, and deliver a complete "
        "industry platform, not a slide deck.",
        space_after=8,
    )

    add_heading(doc, "Why KnightRoad", level=2)
    add_body(
        doc,
        "We build AI platforms that embed intelligence into daily workflows — automated "
        "scheduling, rule-driven safety compliance, template-based project bootstrap, in-browser "
        "building design with IFC4 export, and a roadmap to live camera vision and enterprise "
        "integrations (Autodesk, ERP, LMS, IoT). Our model is repeatable: identify the pain "
        "point, encode domain expertise into intelligent agents and automation, and deliver a "
        "unified platform leaders can market and deploy.",
        space_after=8,
    )

    add_heading(doc, "Ask", level=2)
    add_bullet(doc, "Use this platform as the lead case study for construction and vocational-education prospects.")
    add_bullet(doc, "Position KnightRoad as the partner that ships working AI platforms, not pilots that stall.")
    add_bullet(doc, "Prioritize sales conversations with contractors, vocational institutes, and owner-operators who need one system instead of five.")

    add_page_break(doc)

    # --- FULL BUSINESS CASE ---
    add_heading(doc, "Business Case: Veritas AI Construction Platform", level=1)

    add_heading(doc, "1. Company Overview", level=2)
    add_body(
        doc,
        "KnightRoad Veritas AI Platforms designs and delivers AI-powered business platforms "
        "that address specific industry pain points. Our Veritas product line combines "
        "predictive analytics, workflow automation, and domain-specific intelligence so "
        "organizations can explore and act on their operational data with confidence.",
    )
    add_body(
        doc,
        "Veritas AI Construction Platform is the construction vertical — built to replace "
        "tool sprawl with one integrated environment for project leaders, safety officers, "
        "site foremen, and training programs.",
    )

    add_heading(doc, "2. Problem Statement", level=2)
    add_body(doc, "Construction and vocational training organizations consistently face:")

    problems = [
        ("Slow project startup", "Manual setup of schedules, safety plans, zones, and training assignments takes weeks and invites errors."),
        ("Disconnected systems", "Scheduling, safety, BIM models, assets, and training live in separate tools with no shared context."),
        ("CAD software dependency", "Teams without Revit or AutoCAD licences cannot create or iterate on building designs; early-stage layout work is blocked."),
        ("Compliance exposure", "Safety rules are documented but not actively monitored; alerts are reactive, not tied to the live schedule."),
        ("Workforce readiness gaps", "Workers arrive on site without verified competency; training records are not linked to project roles."),
        ("Limited executive visibility", "Leaders lack a single view of budget, progress, risk, and team status across active jobs."),
        ("BIM data locked away", "3D models sit with designers while field teams work from outdated 2D plans."),
    ]
    for title, desc in problems:
        add_bullet(doc, f" — {desc}", bold_prefix=title + ":")

    add_heading(doc, "3. Solution Overview", level=2)
    add_body(
        doc,
        "Veritas AI Construction Platform is a web-based command center that unifies building "
        "design creation, project setup, live monitoring, resource planning, safety operations, "
        "asset tracking, BIM visualization, and VR workforce training. Users work within a "
        "single project context across every module — no re-entry of data, no context switching, "
        "no dependency on external CAD tools.",
    )

    add_heading(doc, "4. Platform Capabilities", level=2)

    modules = [
        (
            "Executive Dashboard",
            "Real-time project health at a glance: budget vs. spend, progress against target, "
            "recent safety alerts, upcoming tasks, VR training status, team roster, project "
            "documents, and an interactive 3D building model filtered by construction phase.",
        ),
        (
            "Design Studio",
            "In-browser building design environment — draw 2D floor plans storey by storey "
            "(walls, columns, doors, windows, MEP duct and pipe runs), get an instant live 3D "
            "preview with construction-phase filtering, and export a valid IFC4 file for use in "
            "Revit, ArchiCAD, or Blender. Auto-detects and names rooms from wall geometry. "
            "Designs render directly in the dashboard BIM viewer and can be pushed to Speckle "
            "for team sharing and version history. No external CAD licence required.",
        ),
        (
            "New Project Wizard",
            "A guided 10-step flow that takes a blank project to fully operational. Users select "
            "a building type, enter project details, define site zones, assign team members, "
            "review an auto-generated schedule, confirm safety protocols, review VR assignments, "
            "upload documents and BIM files, and publish — with draft save/resume at every step.",
        ),
        (
            "Resource Plan",
            "Operational scheduling workspace combining a searchable Kanban board with a Gantt "
            "chart using critical-path method scheduling, task dependencies, crew leveling, and "
            "budget linkage to completed work.",
        ),
        (
            "Safety Monitor",
            "Site safety operations center with zone status cards, a live alert log (severity, "
            "confidence, acknowledgement workflow), and camera-mapped site zones. Alerts sync "
            "with the executive dashboard so nothing is missed.",
        ),
        (
            "Resourciist (Resource List)",
            "Inventory of personnel, heavy machinery, and materials with availability status, "
            "location tracking, and search/filter across the project.",
        ),
        (
            "VR Training Hub",
            "Role-based VR module assignments (mandatory, recommended, or not required) with "
            "completion tracking. Modules cover crane operation, fall protection, site safety, "
            "steel assembly, forklift operation, infection control, and more.",
        ),
        (
            "BIM Viewer",
            "In-browser 3D visualization of IFC, DWG, DXF, and Design Studio models with "
            "construction-phase filtering — bringing design data to field teams without desktop "
            "CAD software.",
        ),
    ]
    for name, desc in modules:
        add_bullet(doc, f" — {desc}", bold_prefix=name + ":")

    add_heading(doc, "5. Intelligent Automation (AI Agents)", level=2)
    add_body(
        doc,
        "Veritas embeds domain intelligence directly into workflows. These are not generic "
        "chatbots — they are purpose-built automation agents that act on construction "
        "knowledge:",
    )
    agents = [
        ("Project Bootstrap Agent", "Selects building-type templates and auto-populates site zones, task sequences, resource lists, safety rules, and VR training matrices."),
        ("Schedule Agent", "Generates critical-path schedules scaled to team size, project complexity, and working-day calendars, with crew leveling across resource pools."),
        ("Safety Compliance Agent", "Maps OSHA, NFPA, CDC, and institutional safety codes to site zones and generates alerts when schedule or inspection milestones are at risk."),
        ("Training Assignment Agent", "Auto-assigns VR competency modules based on each team member's role at project creation."),
        ("Alert Synthesis Agent", "Derives actionable alerts from schedule state — overdue tasks, behind-schedule work, inspection due dates — and pushes them to dashboard and safety views in real time."),
        ("Parametric Geometry Agent", "Translates simple building parameters (footprint, storeys, bay grid, window ratio, roof type) or a hand-drawn 2D floor plan into a fully tessellated 3D model with IFC-classified elements — no external geometry software involved."),
        ("IFC Authoring Agent", "Converts any Design Studio model (drawn or parametric) into a compliant IFC4 file with a full spatial hierarchy (Project → Site → Building → Storeys) and correctly typed elements, ready for import into Revit, ArchiCAD, or Blender."),
        ("Vision Agent (Roadmap)", "Architecture in place for live camera feeds with AI detection of PPE violations, exclusion zone breaches, and unsafe conditions."),
    ]
    for name, desc in agents:
        add_bullet(doc, f" — {desc}", bold_prefix=name + ":")

    add_body(
        doc,
        "This agent-based architecture is KnightRoad's core capability: we encode industry "
        "expertise into intelligent automation that runs continuously, not on demand.",
        italic=True,
    )

    add_heading(doc, "6. Building Type Templates", level=2)
    add_body(
        doc,
        "Eight pre-configured templates accelerate project startup across project types:",
    )
    templates = [
        "Single-Family Residence",
        "Multi-Family / Apartment",
        "Vocational / Academic (primary education use case)",
        "Office / Commercial",
        "Industrial Warehouse",
        "Healthcare Facility",
        "Retail / Mixed-Use",
        "Infrastructure / Civil",
    ]
    for t in templates:
        add_bullet(doc, t)

    add_heading(doc, "7. Target Market", level=2)
    markets = [
        ("Vocational & technical schools", "Programs that teach construction while running live building projects — dual-purpose teaching and operations."),
        ("Small-to-mid contractors", "Firms that cannot afford enterprise PM suites but need integrated scheduling, safety, and training."),
        ("Owner-operators & institutions", "Schools, hospitals, and municipalities managing capital construction with compliance requirements."),
        ("General contractors (pilot expansion)", "Teams seeking a unified platform before committing to full ERP or Primavera deployments."),
    ]
    for title, desc in markets:
        add_bullet(doc, f" — {desc}", bold_prefix=title + ":")

    add_heading(doc, "8. Competitive Differentiation", level=2)
    diffs = [
        "Purpose-built for construction operations and vocational education — not a generic project management tool repackaged.",
        "Full design-to-delivery lifecycle in one platform: draw the building in-browser, export IFC, initialize the project, manage the schedule, monitor safety, and train the workforce — without leaving the app.",
        "Template-driven project bootstrap reduces setup from weeks to hours.",
        "Safety, schedule, and training alerts share one engine — no duplicate data entry or missed signals.",
        "In-browser BIM creation (Design Studio) eliminates the Revit/AutoCAD dependency for teams that need to sketch, iterate, or teach without expensive CAD licences.",
        "IFC4 export from the Design Studio produces industry-standard files that import into any major BIM tool.",
        "Speckle integration enables cloud-based model versioning and sharing without requiring Autodesk infrastructure.",
        "VR training woven into project onboarding, not bolted on as an afterthought.",
        "Agent-based automation encodes domain expertise — schedules, safety codes, training matrices, geometry generation — without requiring users to prompt an AI chatbot.",
        "Clear expansion path to live IoT cameras, Autodesk Construction Cloud, Moodle LMS, SAP ERP, and enterprise scheduling tools.",
    ]
    for d in diffs:
        add_bullet(doc, d)

    add_heading(doc, "9. Current Maturity & Deployment", level=2)
    add_body(doc, "Platform status as of June 2026:")
    add_bullet(doc, "Version 1.0 deployed as a working web application (Docker-ready, cloud-hosted).")
    add_bullet(doc, "Live client reference: BTVI Vocational School Construction Department, Nassau, Bahamas.")
    add_bullet(doc, "Flagship project: Vocational Center Phase 1 (BSD$1.5M academic block and workshop wings).")
    add_bullet(doc, "Full wizard workflow, CPM scheduling, BIM viewer, multi-module UI, and real-time dashboard refresh operational.")
    add_bullet(doc, "Design Studio operational: in-browser 2D plan drawing, instant 3D preview, IFC4 export, and Speckle push all live.")
    add_bullet(doc, "Enterprise integrations (Autodesk, AWS IoT, SAP, Moodle, Primavera) architected and documented for Phase 2.")

    add_heading(doc, "10. Business Outcomes", level=2)
    outcomes = [
        ("Faster project startup", "Hours instead of weeks to go from blank slate to published, operational project."),
        ("Design without CAD licences", "Teams sketch, iterate, and produce IFC4 models directly in the browser — eliminating Revit/AutoCAD as a prerequisite."),
        ("Reduced compliance risk", "Safety rules mapped to zones with active alert monitoring and acknowledgement trails."),
        ("Better schedule adherence", "Critical-path visibility with overdue and behind-schedule alerts pushed to leaders automatically."),
        ("Workforce readiness", "Role-based VR training assigned at project creation with tracked completion before site work."),
        ("Executive clarity", "One dashboard for budget, progress, safety, tasks, team, documents, and 3D model."),
        ("Lower tool cost", "One platform replaces multiple point solutions for design sketching, scheduling, safety logging, asset tracking, and training records."),
    ]
    for title, desc in outcomes:
        add_bullet(doc, f" — {desc}", bold_prefix=title + ":")

    add_heading(doc, "11. Revenue & Go-to-Market Positioning", level=2)
    add_body(
        doc,
        "Market Veritas AI Construction Platform as a subscription SaaS offering tiered by "
        "active projects and user seats. Lead with the BTVI case study — a vocational "
        "institute running a live capital project on the platform. Emphasize KnightRoad's "
        "delivery model: we ship working platforms with embedded AI agents, not proof-of-concept "
        "demos that never reach production.",
    )
    add_body(doc, "Suggested messaging pillars for sales and marketing:")
    pillars = [
        '"One platform. One project context. Zero tool sprawl."',
        '"From blank site to operational project in a single afternoon."',
        '"Design your building in the browser. Export IFC. No CAD licence needed."',
        '"Safety, schedule, and training — connected, not siloed."',
        '"AI that works in the background so your team works on the build."',
        '"Built by KnightRoad Veritas AI Platforms — predictive analytics you can deploy today."',
    ]
    for p in pillars:
        add_bullet(doc, p)

    add_heading(doc, "12. Roadmap Highlights", level=2)
    roadmap = [
        "Live camera integration with AI vision for PPE and zone breach detection",
        "Autodesk Construction Cloud BIM sync",
        "Moodle LMS integration for VR module delivery",
        "SAP ERP asset management connector",
        "Enterprise scheduling import/export (Primavera, MS Project)",
        "Design Studio: structural analysis hints and clash detection",
        "Design Studio: collaborative multi-user floor plan editing",
        "Multi-tenant SaaS for regional vocational networks and contractor associations",
        "Mobile field app for site foremen and safety officers",
    ]
    for r in roadmap:
        add_bullet(doc, r)

    add_heading(doc, "13. Recommendation", level=2)
    add_body(
        doc,
        "Approve Veritas AI Construction Platform as KnightRoad's lead construction vertical "
        "offering. The product is demo-ready, client-proven, and architected for scale. "
        "Direct sales and marketing effort toward vocational institutes and small-to-mid "
        "contractors in the Caribbean and North America, using the BTVI deployment as the "
        "anchor reference.",
    )
    add_body(
        doc,
        "KnightRoad Veritas AI Platforms is positioned to repeat this model across other "
        "industries — identify the pain point, encode expertise into intelligent agents, "
        "deliver a unified platform, and prove value with a live client before scaling.",
        bold=True,
    )

    add_horizontal_rule(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run(
        "KnightRoad Veritas AI Platforms  |  Veritas AI Construction Platform  |  Confidential"
    )
    set_run_font(f_run, size_pt=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
