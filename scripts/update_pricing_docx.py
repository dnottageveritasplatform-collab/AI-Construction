# -*- coding: utf-8 -*-
"""Update PRICING STRATEGIES docx for Design Studio and revised price points."""
import datetime
import os
import shutil
from docx import Document
from docx.oxml.ns import qn

SRC = r"j:\My Drive\Vertias_AI_Contruction\PRICING STRATEGIES (Veritas AI Construction).docx"
LOCAL = r"c:\Users\Dominic Nottage\aethel-protocol\Veritas_AI_Construction\docs\PRICING STRATEGIES (Veritas AI Construction).docx"


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)


def line(doc, segments):
    p = doc.add_paragraph()
    for text, bold in segments:
        r = p.add_run(text)
        r.bold = bold
    return p


def blank(doc):
    doc.add_paragraph()


def header(doc, text):
    return line(doc, [(text, True)])


def item(doc, name, desc, price):
    line(doc, [(name + " \u2014 ", True), (desc + " ", False), ("(" + price + ")", True)])


def bullet(doc, label, body):
    line(doc, [(label + ": ", True), (body, False)])


def build_document():
    doc = Document(SRC)
    clear_body(doc)

    header(doc, "PRICING STRATEGIES:")
    blank(doc)
    line(doc, [(
        "Platform update (Design Studio): Veritas now includes an in-browser Design Studio "
        "workspace \u2014 2D floor-plan authoring (walls, doors, windows, columns, MEP), "
        "multi-storey editing, room detection, instant 3D preview, IFC4 export, and Speckle "
        "cloud push. Saved designs feed the dashboard BIM viewer automatically, closing the "
        "design-to-delivery loop without Revit or AutoCAD. This materially increases platform "
        "value for vocational programs, owner-operators, and SMB contractors.",
        False,
    )])
    blank(doc)

    # --- Option A ---
    header(doc, "Option A \u2014 Monthly licensing (you keep ownership)")
    blank(doc)
    header(doc, "Per-seat tiers:")
    bullet(doc, "Starter", "~$65/user/mo \u2014 dashboard, project/doc management, team roster, KPIs")
    bullet(doc, "Professional", "~$179/user/mo \u2014 + BIM/IFC viewer, CAD/DWG processing, Gantt/resource plan, Design Studio (2D editor + 3D preview + save)")
    bullet(doc, "Enterprise", "~$349/user/mo \u2014 + full Design Studio (IFC4 export, Speckle push, templates), safety monitoring, VR training, SSO, integrations")
    blank(doc)
    line(doc, [("Annual prepay discount: ", True), ("~20% off monthly rates when billed annually.", False)])
    blank(doc)
    header(doc, "Per-organization (better fit for schools/contractors):")
    bullet(doc, "Single site", "$1,200\u2013$2,250/mo \u2014 one active project; includes Design Studio seats for design team")
    bullet(doc, "Multi-project firm", "$3,750\u2013$9,000/mo \u2014 several concurrent builds; unlimited Design Studio authoring within org")
    bullet(doc, "Institution / district", "$32k\u2013$95k/yr \u2014 vocational schools and training programs (Design Studio is the primary differentiator)")
    blank(doc)
    header(doc, "Add-ons:")
    bullet(doc, "Onboarding / implementation", "$5k\u2013$35k \u2014 data migration, integration wiring, Design Studio instructor setup")
    bullet(doc, "Training (incl. Design Studio + VR)", "$2,500\u2013$15k \u2014 per cohort or institution")
    bullet(doc, "Priority support / SLA", "+15\u201325% of license")
    bullet(doc, "White-label / OEM", "Custom \u2014 partner resells under their brand")
    blank(doc)
    line(doc, [("Pros: ", True), (
        "Recurring/compounding revenue; you keep all upside; Design Studio creates stickier "
        "accounts and higher seat counts; builds ARR that makes a future sale much bigger.",
        False,
    )])
    line(doc, [("Cons: ", True), (
        "Slow cash ramp; ongoing dev/support/sales burden; must replace mock data with live "
        "integrations; churn + liability risk on safety features.",
        False,
    )])
    blank(doc)
    blank(doc)

    # --- Option B ---
    header(doc, "Option B \u2014 Sell the IP (one payday)")
    blank(doc)
    line(doc, [(
        "Inbound buyer interest exists. Design Studio (IFC authoring + Speckle + in-browser "
        "CAD-lite) adds a distinct IP asset beyond the original operations platform. Run a "
        "light competitive process so multiple interested parties bid against each other.",
        False,
    )])
    blank(doc)
    header(doc, "Realistic ranges by stage:")
    bullet(doc, "As-is (pre-revenue, 1 client, Design Studio live)", "$350k\u2013$550k \u2014 code + concept + authoring IP; SME-validated demand")
    bullet(doc, "BTVI live + early pilots", "$475k\u2013$950k \u2014 proof of real-world deployment")
    bullet(doc, "~$200k ARR with real integrations", "$1.0M\u2013$2.2M+ \u2014 multiple of proven recurring revenue")
    blank(doc)
    line(doc, [("Pros: ", True), (
        "Immediate liquidity; no ongoing burden; clean exit; offloads safety liability; "
        "Design Studio IP is hard to replicate quickly.",
        False,
    )])
    line(doc, [("Cons: ", True), (
        "Still likely leaving money on the table vs. licensing; you forfeit all future upside; "
        "non-competes can lock you out of the space.",
        False,
    )])
    blank(doc)
    blank(doc)

    # --- Founder engagement ---
    header(doc, "Founder Development Engagement (Post-Sale)")
    line(doc, [(
        "Separate from the IP purchase price. The buyer acquires the IP and retains the creator "
        "(Dominic R. Nottage) on a dedicated basis to build out the roadmap. Charge IP and "
        "development as two distinct line items.",
        False,
    )])
    blank(doc)
    header(doc, "Recommended Engagement \u2014 Dedicated Monthly Retainer")
    bullet(doc, "Basis", "~160 hrs/month at $110\u2013$185/hr (original author, PMP \u00b7 CSM, construction-domain + Design Studio expertise)")
    bullet(doc, "Recommended", "$26k\u2013$28k/mo (~$312k\u2013$336k/yr) \u2014 blended ~$165/hr, founder/domain premium")
    bullet(doc, "Lower bound", "~$18k/mo (~$216k/yr) \u2014 $110/hr equivalent")
    bullet(doc, "Upper bound", "~$32k/mo (~$384k/yr) \u2014 $185/hr equivalent")
    blank(doc)
    header(doc, "Engagement Terms (protect against open-ended scope)")
    bullet(doc, "Minimum term", "12 months, then month-to-month or renew")
    bullet(doc, "Defined capacity", "Up to agreed hrs/month; beyond that bills at $185/hr or rolls forward")
    bullet(doc, "Change-order clause", "New scope requires SOW or written approval")
    bullet(doc, "Annual rate review", "5\u20138%/yr or re-benchmark to market")
    bullet(doc, "New-work IP", "Assigns to buyer (they own it) \u2014 rate reflects premium end")
    blank(doc)
    header(doc, "Roadmap Menu \u2014 \u00c0 La Carte Reference Prices")
    line(doc, [(
        "Under the retainer these are deliverables your dedicated time produces. Ranges also "
        "serve as fall-back pricing for individual features. Low end = tight scope; high end = "
        "polish + testing.",
        False,
    )])
    blank(doc)
    header(doc, "Design Studio enhancements (new)")
    item(doc, "Collaborative multi-user editing", "real-time co-editing on floor plans", "$35k\u2013$75k")
    item(doc, "Clash detection (Design Studio + viewer)", "flag spatial conflicts before export", "$28k\u2013$55k")
    item(doc, "Structural / code hints", "basic span/load checks and advisory flags", "$22k\u2013$50k")
    item(doc, "Expanded template library", "sector-specific starter plans beyond BT-01\u2013BT-08", "$12k\u2013$28k")
    item(doc, "Design Studio mobile / tablet", "field sketching and markup on site", "$30k\u2013$65k")
    blank(doc)
    header(doc, "Integrations (replace stubbed connectors)")
    item(doc, "Autodesk Construction Cloud / BIM 360", "live BIM instead of manual IFC uploads", "$15k\u2013$30k")
    item(doc, "IoT safety pipeline (MQTT / AWS IoT)", "real-time alerts from sensors/cameras", "$15k\u2013$32k")
    item(doc, "Scheduling sync (MS Project / Primavera)", "two-way task / Gantt sync", "$15k\u2013$30k")
    item(doc, "LMS integration (Moodle)", "course / progress tracking for VR training", "$10k\u2013$22k")
    item(doc, "ERP / asset integration (SAP, QuickBooks)", "live inventory & cost data", "$15k\u2013$35k")
    blank(doc)
    header(doc, "AI differentiation")
    item(doc, "Computer-vision PPE / safety detection", "auto-flag hardhat/vest/zone violations", "$30k\u2013$70k")
    item(doc, "Schedule & cost risk prediction", "ML surfaces at-risk tasks/budget early", "$22k\u2013$52k")
    item(doc, "Progress-from-photos vs BIM", "track % complete from site photos", "$30k\u2013$70k")
    item(doc, "AI document / RFI assistant (LLM + RAG)", "Q&A over specs, drawings, submittals", "$22k\u2013$48k")
    blank(doc)
    header(doc, "Reach & experience")
    item(doc, "Mobile field app (PWA or native)", "foreman / crew access on site", "$30k\u2013$80k")
    item(doc, "Advanced BIM viewer", "measurements, markups, phase compare", "$22k\u2013$52k")
    item(doc, "VR training authoring + headset support", "build courses, run on real VR hardware", "$30k\u2013$62k")
    item(doc, "Reporting & analytics + exports", "dashboards, PDF/Excel, audit logs", "$10k\u2013$22k")
    blank(doc)
    line(doc, [("Note: ", True), (
        "Phase 1 productionization (real database, auth / RBAC, cloud deployment & security "
        "hardening) is excluded \u2014 the creator is completing those before the product is "
        "presented for sale.",
        False,
    )])
    blank(doc)
    header(doc, "Suggested Pitch to the Buyer")
    line(doc, [(
        "\u201cYou acquire the IP including Design Studio. I stay on dedicated at "
        "$26k\u2013$28k/mo (12-month minimum) to execute the roadmap \u2014 Design Studio "
        "enhancements, integrations, AI features, and field/VR reach \u2014 with new scope "
        "handled via change orders. \u00c0 la carte pricing is available if you\u2019d prefer "
        "specific features instead of a retainer.\u201d",
        False,
    )])

    return doc


def main():
    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = SRC.replace(".docx", f" BACKUP {stamp}.docx")
    try:
        shutil.copy2(SRC, backup)
        print("Backup:", os.path.basename(backup))
    except OSError as exc:
        print("Backup skipped (source locked?):", exc)

    doc = build_document()

    os.makedirs(os.path.dirname(LOCAL), exist_ok=True)
    doc.save(LOCAL)
    print("Saved locally:", LOCAL)

    try:
        doc.save(SRC)
        print("Updated Google Drive copy:", SRC)
    except PermissionError:
        print("Google Drive copy locked — close the file in Word, then run:")
        print(f'  Copy-Item -LiteralPath "{LOCAL}" -Destination "{SRC}" -Force')
    print("Paragraphs:", len(doc.paragraphs))


if __name__ == "__main__":
    main()
