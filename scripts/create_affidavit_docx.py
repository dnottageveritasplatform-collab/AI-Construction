# -*- coding: utf-8 -*-
"""Create the Veritas affidavit as a Word document."""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

OUT = r"j:\My Drive\Vertias_AI_Contruction\Affidavit - KnightRoad Veritas AI Construction Platform.docx"
LOCAL = r"c:\Users\Dominic Nottage\aethel-protocol\Veritas_AI_Construction\docs\Affidavit - KnightRoad Veritas AI Construction Platform.docx"


def centered(doc, text, bold=False, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def body(doc, text, space_after=12, first_line_indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    return p


def numbered(doc, text, space_after=12):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def signature_block(doc, name, space_before=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(name)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(18)
    lr = line.add_run("_" * 40)
    lr.font.size = Pt(12)
    lr.font.name = "Times New Roman"


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    centered(doc, "COMMONWEALTH OF THE BAHAMAS", bold=True, space_after=4)
    centered(doc, "NASSAU N. P BAHAMAS", bold=True, space_after=4)
    centered(doc, "AFFIDAVIT", bold=True, space_after=12)
    centered(doc, "REFERENCE: NIGHTROAD VERITAS AI CONSTRUCTION PLATFORM", bold=True, space_after=24)

    body(
        doc,
        "I Dominic Nottage age 47 of Sandilands Village, owner of the above business, "
        "being in my sound mind, stand before Bishop Philemon R. Wilson a Justice of the "
        "Peace of the above Island Nassau N.P Bahamas and state the following:",
        space_after=18,
    )

    numbered(
        doc,
        "That I have entered into agreement with Shorn Gibson age 55 of Yamacraw Shores "
        "and Shannon Williams age 32 of Bitter Road Street Nassau N.P Bahamas, both being "
        "of sound mind",
    )
    numbered(
        doc,
        "We agree that Shorn Gibson & Shannon Williams would be subject matters experts "
        "to the above business",
    )
    numbered(
        doc,
        "That that shares from the gross revenue would be shared quarterly in the "
        "following percentage: Dominic Nottage owner of Business 49% Shorn Gibson 26% & "
        "Shannon Williams 25%",
    )
    numbered(
        doc,
        "I further swear that this business is registered in the above name in the "
        "Bahamas and the United States of America, with all rights of ownership to "
        "Dominic Nottage.",
    )
    numbered(
        doc,
        "We sign this agreement this 22nd day of July 2026 in the year of our Lord",
        space_after=24,
    )

    signature_block(doc, "Dominic Nottage", space_before=12)
    signature_block(doc, "Shorn Gibson")
    signature_block(doc, "Shannon Williams")

    body(
        doc,
        "Before Me Bishop Philemon R Wilson Justice of the Peace Nassau N.P Bahamas",
        space_after=0,
        first_line_indent=0,
    )

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(LOCAL), exist_ok=True)
    doc.save(LOCAL)
    print("Saved locally:", LOCAL)
    try:
        doc.save(OUT)
        print("Saved to Drive:", OUT)
    except PermissionError:
        print("Drive path locked; local copy is ready.")
        print(f'Copy with: Copy-Item -LiteralPath "{LOCAL}" -Destination "{OUT}" -Force')


if __name__ == "__main__":
    main()
