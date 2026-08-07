# -*- coding: utf-8 -*-
"""
Generate IP ownership letter and development cost breakdown documents.
Run: py scripts/generate_ip_ownership_letter.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCS = PROJECT_ROOT / "docs"
LETTER_PATH = DOCS / "Letter - IP Ownership Clarification (Veritas AI Construction).docx"
COST_PATH = DOCS / "Development Cost Breakdown - Veritas AI Construction Platform.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def body(doc, text, bold=False, italic=False, space_after=8, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    font(run, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        font(r1, bold=True)
        r2 = p.add_run(text)
        font(r2)
    else:
        run = p.add_run(text)
        font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            font(run, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                font(run, size=10)
    doc.add_paragraph()
    return t


def build_letter():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    today = date.today().strftime("%B %d, %Y")

    body(doc, today, space_after=18)
    body(doc, "Shorn Gibson", space_after=2)
    body(doc, "Yamacraw Shores", space_after=2)
    body(doc, "Nassau, N.P., The Bahamas", space_after=18)

    body(doc, "RE: Intellectual Property Ownership — KnightRoad Veritas AI Construction Platform", bold=True, space_after=18)

    body(doc, "Dear Shorn,", space_after=12)

    body(
        doc,
        "I am writing to address recent statements suggesting that you have ownership of, or "
        "primarily originated, the KnightRoad Veritas AI Construction Platform (\"the Platform\"). "
        "Those statements are inconsistent with the facts, with the written agreements we signed, "
        "and with how the product was actually conceived, designed, built, funded, and deployed. "
        "This letter sets out the record clearly so there is no further misunderstanding.",
        space_after=12,
    )

    heading(doc, "1. How the Platform Was Conceived", level=2)
    body(
        doc,
        "The Platform did not begin as your original concept. I developed and owned a separate "
        "application — the Global Monitor platform (Gemini Code workspace) — before any "
        "construction-specific product existed. During a discussion, you observed that a similar "
        "type of system could be applied to the construction industry. That was a directional "
        "suggestion about market vertical, not the creation of the Platform itself.",
        space_after=8,
    )
    body(
        doc,
        "Applying an existing product architecture to a new industry is a standard product "
        "strategy. It does not transfer ownership of the resulting application, its design, "
        "its requirements, its source code, or its intellectual property to the person who "
        "suggested the vertical. The Platform was built by adapting my pre-existing work, "
        "my design framework, and my development process — not by implementing an idea you "
        "documented, designed, or coded.",
        space_after=12,
    )

    heading(doc, "2. What I Created and Delivered", level=2)
    body(doc, "I personally performed the substantive work required to bring the Platform into existence, including:", space_after=8)
    bullet(doc, "Product vision, information architecture, and UI/UX design across all modules.")
    bullet(doc, "Requirements analysis, use-case documentation, and technical specifications.")
    bullet(doc, "Complete source-code development (backend API, frontend modules, integrations).")
    bullet(doc, "Design Studio — in-browser 2D floor-plan authoring, 3D preview, IFC4 export, and Speckle collaboration.")
    bullet(doc, "BIM/IFC/DWG viewing pipeline, project wizard, dashboard, safety monitor, resource planning, VR training hub, and related modules.")
    bullet(doc, "Testing, debugging, repository creation on GitHub, and live deployment on Render.")
    bullet(doc, "Business documentation, design documents, user guides, and demo materials.")
    bullet(doc, "Copyright registration of the source code and design under my company, KnightRoad Veritas AI Platforms.")
    body(
        doc,
        "You did not write requirements documents, technical specifications, or source code. "
        "You did not fund development, pay for AI tooling tokens, or pay hosting costs during "
        "the build phase. The Platform exists because of work I performed, at my expense, "
        "using skills and assets I already possessed.",
        space_after=12,
    )

    heading(doc, "3. Your Stated Contributions", level=2)
    body(
        doc,
        "You have participated in the SDLC process at a high level by offering general feedback "
        "such as \"it's gotta pop,\" \"make it easy to use,\" and \"it's got to stand out.\" "
        "That type of subjective input is common in early product discussions and can be valuable "
        "as user feedback — but it is not equivalent to:",
        space_after=8,
    )
    bullet(doc, "Authorship or co-authorship of the software.")
    bullet(doc, "Ownership of intellectual property.")
    bullet(doc, "Creation of the product concept, design system, or technical architecture.")
    bullet(doc, "Investment in or payment for development.")
    body(
        doc,
        "Industry-standard product development treats such feedback as stakeholder input, not "
        "as a basis for co-ownership of the underlying IP.",
        space_after=12,
    )

    heading(doc, "4. What Our Signed Agreements Say", level=2)
    body(
        doc,
        "Our Revenue Share License Agreement (drafted April 22, 2026; amended May 25, 2026) and "
        "the Affidavit signed July 22, 2026, record the arrangement explicitly:",
        space_after=8,
    )
    bullet(doc, "I (Dominic Nottage) retain 100% ownership of all intellectual property and source code.", bold_prefix="IP ownership: ")
    bullet(doc, "Any improvements or enhancements developed during the partnership remain my exclusive property, regardless of who suggested them.", bold_prefix="Improvement clause: ")
    bullet(doc, "You and Shannon Williams participate as partners entitled to defined percentages of net revenue (26% and 25% respectively) — not as co-owners of the Platform.", bold_prefix="Revenue share only: ")
    bullet(doc, "The business is registered with all rights of ownership to me in The Bahamas and the United States.", bold_prefix="Registration: ")
    body(
        doc,
        "Those documents do not grant you equity in the software, title to the codebase, or "
        "the right to represent yourself as the owner or originator of the Platform. Your role "
        "under the agreement is as a revenue-share partner and subject-matter expert — not as "
        "an IP owner.",
        space_after=12,
    )

    heading(doc, "5. Why \"It Was My Idea\" Does Not Create Ownership", level=2)
    body(
        doc,
        "Ideas alone are not protectable assets in the way that executed software, documented "
        "design, and registered copyright are. Courts and commercial practice consistently "
        "recognize that ownership follows creation: who designed it, who built it, who paid "
        "for it, and who the contracts say owns it.",
        space_after=8,
    )
    body(
        doc,
        "Suggesting that a product should exist in a particular industry is not the same as "
        "inventing, designing, specifying, coding, testing, deploying, and copyrighting that "
        "product. If vertical suggestions created ownership, every advisor, customer, and "
        "casual observer who said \"you should build X for Y industry\" would hold a claim to "
        "the resulting software — which is not how software businesses operate.",
        space_after=12,
    )

    heading(doc, "6. Financial and Operational Facts", level=2)
    bullet(doc, "You have not paid for development labor, AI tooling, or hosting during the build phase.")
    bullet(doc, "I created and maintain the GitHub repository and Render deployment.")
    bullet(doc, "I registered copyright under my company.")
    bullet(doc, "The Platform is deployed and operational as a direct result of my work product.")
    body(
        doc,
        "A separate development-cost analysis (attached or available upon request) estimates "
        "that reproducing this Platform with a professional team would cost approximately "
        "$400,000 to $750,000 USD and require roughly 2,500 to 4,500 engineering hours. "
        "That investment was borne by me, not by you.",
        space_after=12,
    )

    heading(doc, "7. Required Understanding Going Forward", level=2)
    body(doc, "To avoid further dispute, the following must be understood and respected:", space_after=8)
    bullet(doc, "I am the sole owner of the Platform's intellectual property and source code.")
    bullet(doc, "You are a revenue-share partner under the signed agreement, not an owner of the software.")
    bullet(doc, "You must not represent to customers, investors, or third parties that you own, authored, or originated the Platform.")
    bullet(doc, "Feature suggestions remain welcome as partner input; they do not convert into IP ownership under our agreement.")
    body(
        doc,
        "I value the domain expertise you bring as a construction subject-matter expert. That "
        "expertise is why the revenue-share structure exists. It is not a substitute for "
        "authorship, development investment, or IP title.",
        space_after=12,
    )

    heading(doc, "8. Conclusion", level=2)
    body(
        doc,
        "The Platform is my work product. It was built on my prior application architecture, "
        "designed and coded by me, deployed by me, copyrighted by my company, and contractually "
        "documented as my exclusive property. Your suggestion to apply a monitoring-style concept "
        "to construction may have helped identify a market — and for that you are compensated "
        "through revenue share — but it does not make you the owner of the Platform.",
        space_after=12,
    )
    body(
        doc,
        "If you continue to assert ownership or misrepresent your role to third parties, I "
        "will need to enforce the written agreements and protect the IP accordingly. I prefer "
        "to resolve this cooperatively and focus on growing the business within the structure "
        "we already signed.",
        space_after=24,
    )

    body(doc, "Respectfully,", space_after=36)
    body(doc, "Dominic R. Nottage, PMP · CSM", bold=True, space_after=4)
    body(doc, "Founder & Owner — KnightRoad Veritas AI Platforms", space_after=4)
    body(doc, "Licensor — Veritas AI Construction Platform", space_after=18)

    body(
        doc,
        "This letter is a factual statement of record based on signed agreements and development "
        "history. It is not legal advice. Either party may consult independent counsel.",
        italic=True,
    )
    p = doc.paragraphs[-1]
    for run in p.runs:
        font(run, size=9, italic=True)

    return doc


def build_cost_breakdown():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(title.add_run("Development Cost Breakdown"), size=16, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        subtitle.add_run("KnightRoad Veritas AI Construction Platform — v1.0"),
        size=12,
        italic=True,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(meta.add_run(f"Prepared: {date.today().strftime('%B %d, %Y')}  |  Confidential"), size=9, italic=True)

    doc.add_paragraph()

    heading(doc, "Purpose", level=2)
    body(
        doc,
        "This document estimates what it would cost a professional software team to build the "
        "Veritas AI Construction Platform from scratch to its current v1.0 scope — including "
        "Design Studio, BIM pipeline, operational modules, documentation, and deployment. "
        "Figures are illustrative planning estimates based on module scope and prevailing U.S./"
        "remote contractor rates for senior product engineering teams.",
    )

    heading(doc, "Scope Summary (Current v1.0)", level=2)
    modules = [
        "Executive Dashboard with real-time SSE updates, KPI tiles, 3D phase viewer",
        "Design Studio — 2D floor-plan editor, multi-storey, rooms, MEP, 3D preview, IFC4 export, Speckle push",
        "New Project Wizard — 10-step guided setup with 8 building templates",
        "BIM/IFC/DWG import and geometry pipeline (phase-aware viewer)",
        "Resource Plan — Kanban + critical-path Gantt scheduling",
        "Safety Monitor — zone alerts, acknowledgement workflow, live feed",
        "Resourciist — personnel, equipment, and materials inventory",
        "VR Training Hub — role-based module assignments and launch",
        "Authentication, project store, REST API layer (14 API modules)",
        "Documentation — design document, business case, user guides, demo assets",
        "DevOps — GitHub repository, Render live hosting",
    ]
    for m in modules:
        bullet(doc, m)

    heading(doc, "Estimated Team Composition", level=2)
    table(
        doc,
        ["Role", "Typical Rate (USD/hr)", "Involvement"],
        [
            ["Senior Full-Stack Engineer (lead)", "$150 – $200", "Architecture, backend, integrations"],
            ["Mid-Level Full-Stack Engineer", "$100 – $130", "Frontend modules, API wiring"],
            ["UI/UX Designer", "$100 – $150", "Dashboard, wizard, Design Studio UX"],
            ["Business Analyst / Product Owner", "$100 – $140", "Requirements, use cases, acceptance criteria"],
            ["QA Engineer", "$80 – $110", "Functional, regression, cross-browser testing"],
            ["DevOps Engineer", "$120 – $160", "CI/CD, cloud deployment, monitoring"],
            ["Technical Writer", "$80 – $100", "User guides, design docs"],
            ["Project Manager", "$120 – $150", "Schedule, coordination, delivery"],
        ],
    )

    heading(doc, "Module-by-Module Build Estimate", level=2)
    body(doc, "Hours and cost ranges assume building to production-ready v1.0 quality, not a prototype.", space_after=8)
    rows = [
        ["Platform foundation (Flask app, auth, config, project store)", "120 – 200", "$18k – $40k"],
        ["Executive Dashboard + real-time SSE", "200 – 350", "$30k – $70k"],
        ["New Project Wizard (10 steps, templates, draft/resume)", "250 – 400", "$38k – $80k"],
        ["Design Studio (2D editor, 3D, rooms, storeys, save)", "400 – 700", "$60k – $140k"],
        ["IFC4 authoring & export pipeline", "120 – 200", "$18k – $40k"],
        ["BIM/IFC/DWG import & phase-aware 3D viewer", "350 – 550", "$53k – $110k"],
        ["Speckle cloud integration (push/pull)", "80 – 150", "$12k – $30k"],
        ["Resource Plan (Kanban + CPM Gantt)", "200 – 350", "$30k – $70k"],
        ["Safety Monitor module", "120 – 200", "$18k – $40k"],
        ["Resourciist (asset inventory)", "80 – 120", "$12k – $24k"],
        ["VR Training Hub", "100 – 180", "$15k – $36k"],
        ["UI/UX design (cross-module)", "200 – 350", "$30k – $70k"],
        ["Requirements, use cases, tech specs", "120 – 200", "$18k – $40k"],
        ["QA & test cycles", "150 – 250", "$23k – $50k"],
        ["Documentation, guides, demo materials", "100 – 180", "$15k – $36k"],
        ["DevOps, deployment, hosting setup", "60 – 120", "$9k – $24k"],
        ["Project management & coordination", "150 – 250", "$23k – $50k"],
    ]
    table(doc, ["Module / Workstream", "Estimated Hours", "Cost Range (USD)"], rows)

    heading(doc, "Totals", level=2)
    table(
        doc,
        ["Scenario", "Total Hours", "Estimated Cost (USD)", "Typical Calendar Time"],
        [
            ["Lean team (blended ~$140/hr)", "2,500 – 3,200", "$350k – $450k", "7 – 9 months"],
            ["Standard agency build (~$165/hr)", "3,000 – 3,800", "$495k – $627k", "8 – 10 months"],
            ["Full team + polish (~$175/hr)", "3,500 – 4,500", "$613k – $788k", "9 – 12 months"],
        ],
    )

    heading(doc, "Additional Costs Not Included Above", level=2)
    bullet(doc, "Third-party SaaS (Render hosting, Speckle, domain, SSL) — ~$100–$500/mo ongoing")
    bullet(doc, "AI/LLM API tokens during development — variable, typically $500–$5,000+ depending on features built")
    bullet(doc, "Software licences (if any commercial BIM/CAD SDKs used)")
    bullet(doc, "Legal (contracts, copyright registration, IP assignment agreements) — $2k–$10k")
    bullet(doc, "Post-launch maintenance — typically 15–20% of build cost annually")

    heading(doc, "What This Demonstrates", level=2)
    body(
        doc,
        "Reproducing the Platform with a professional team would require roughly $400,000 to "
        "$750,000 USD and the better part of a year — before accounting for the founder's "
        "prior Global Monitor codebase, domain expertise (PMP · CSM), and pre-existing design "
        "patterns that accelerated delivery. The actual product includes ~8,600 lines of "
        "backend Python across 14 API modules, ~5,800 lines of HTML templates (including the "
        "Design Studio workspace), and ~7,400 lines of project-authored JavaScript — plus "
        "design documents, business case materials, user guides, and live deployment infrastructure.",
    )
    body(
        doc,
        "This analysis supports the position that the Platform is a substantial, founder-built "
        "software asset — not a casual idea that could be replicated without significant "
        "investment, skill, and execution.",
        italic=True,
    )

    return doc


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    build_letter().save(LETTER_PATH)
    build_cost_breakdown().save(COST_PATH)
    print("Created:")
    print(" ", LETTER_PATH)
    print(" ", COST_PATH)

    drive = Path(r"j:\My Drive\Vertias_AI_Contruction")
    if drive.exists():
        import shutil
        for src in (LETTER_PATH, COST_PATH):
            shutil.copy2(src, drive / src.name)
            print("Copied to Drive:", drive / src.name)


if __name__ == "__main__":
    main()
